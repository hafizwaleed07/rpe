# Research Paper Evaluator (Enhanced with Abstract Findings & Smarter Journal Extraction)

import streamlit as st
import docx
import re
from datetime import datetime
from collections import Counter
import pandas as pd
from io import BytesIO
from docx import Document
from fpdf import FPDF

# Extract text from .docx

def extract_text_from_docx(file):
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs])
# Extract text from PDF using PyPDF2
from PyPDF2 import PdfReader

def extract_text_from_pdf(file):
    pdf = PdfReader(file)
    text = ""
    for page in pdf.pages:
        text += page.extract_text() or ""
    return text
# Detect methodology keywords
def find_keywords(text, keywords):
    return [k for k in keywords if k.lower() in text.lower()]

# Smarter journal name extraction from references
def extract_journal_names_dynamic(text):
    # Extract lines likely from references (heuristic)
    references = re.findall(r'\.\s+([^.\n]{5,100}),\s+\d+\(\d+\)', text)
    journal_candidates = [j.strip() for j in references if j and not j.lower().startswith("doi")]
    cleaned = list(set(journal_candidates))
    return cleaned if cleaned else ["Not Found"]

# Extract abstract section only
def extract_abstract_section(text):
    match = re.search(r'(?i)abstract(.*?)(?=\n\s*[\d]*\s*(introduction|1\.|I\.))', text, re.DOTALL)
    return match.group(1).strip() if match else ""

# Extract findings from abstract section
def extract_key_findings(text):
    abstract = extract_abstract_section(text)
    indicators = [
        "findings reveal", "this study shows", "results indicate",
        "we find that", "empirical results suggest", "the research contributes",
        "reveals that", "confirms that", "our findings", "offers", "draws attention"
    ]
    sentences = re.split(r'(?<=[.!?])\s+', abstract)
    matches = [s.strip() for s in sentences if any(ind in s.lower() for ind in indicators)]
    return matches[:5] if matches else ["Not clearly mentioned"]

# Detect data type
def detect_data_type(text):
    if any(kw in text.lower() for kw in ["secondary data", "archival", "panel data", "financial statements", "firm year observations", "annual reports"]):
        return "Secondary"
    elif any(kw in text.lower() for kw in ["survey", "interview", "questionnaire", "respondents", "participants", "n="]):
        return "Primary"
    return "Not Clear"

# Count recent reference years
def find_recent_years(text):
    current_year = datetime.now().year
    years = re.findall(r'\b(19[7-9]\d|20[0-2]\d|203[0-5])\b', text)
    years = [int(y) for y in years if current_year - int(y) <= 5]
    return dict(Counter(years))

# AI
def detect_ai_like_writing(text):
    ai_phrases = [
        "this study aims to", "it is important to note that",
        "the results indicate that", "the findings suggest that",
        "in conclusion", "this paper highlights"
    ]
    count = sum(text.lower().count(phrase) for phrase in ai_phrases)
    if count >= 5:
        return "⚠️ Writing style appears partially AI-generated. Consider rephrasing or humanizing some sections."
    return "✅ Writing style seems natural."
# Word export
def generate_docx_report(summary):
    doc = Document()
    doc.add_heading('Research Paper Evaluation Summary', 0)
    for section, value in summary.items():
        doc.add_heading(section, level=2)
        if isinstance(value, list):
            for item in value:
                doc.add_paragraph(str(item), style='List Bullet')
        elif isinstance(value, dict):
            for k, v in value.items():
                doc.add_paragraph(f"{k}: {v}", style='List Bullet')
        else:
            doc.add_paragraph(str(value))
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# PDF export
# Generate PDF report (FIXED for Streamlit Download)
def generate_pdf_report(summary):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, "Research Paper Evaluation Summary", ln=True, align='C')
    pdf.ln(10)

    for section, value in summary.items():
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 10, section.upper(), ln=True)
        pdf.set_font("Arial", "", 11)

        if isinstance(value, list):
            if not value:
                pdf.multi_cell(0, 8, "- Not Found")
            else:
                for item in value:
                    pdf.multi_cell(0, 8, f"- {item}")
        elif isinstance(value, dict):
            if not value:
                pdf.multi_cell(0, 8, "- None")
            else:
                for k, v in value.items():
                    pdf.multi_cell(0, 8, f"- {k}: {v}")
        else:
            pdf.multi_cell(0, 8, str(value))
        pdf.ln(5)

    # Output PDF as string and encode to bytes
    pdf_output = pdf.output(dest='S').encode('latin1')
    return BytesIO(pdf_output)

# Streamlit app setup
st.set_page_config(page_title="Research Paper Evaluator", layout="centered")
st.title("📄 Research Paper Evaluator (Enhanced)")

uploaded_file = st.file_uploader("Upload a .docx research paper", type=["docx", "pdf"])

if uploaded_file:
    with st.spinner("Analyzing document..."):
        file_type = uploaded_file.name.split('.')[-1].lower()

        if file_type == "docx":
            text = extract_text_from_docx(uploaded_file)
        elif file_type == "pdf":
            text = extract_text_from_pdf(uploaded_file)
        else:
            st.error("❌ Unsupported file type. Please upload a .docx or .pdf file.")
            st.stop()

        methodology = find_keywords(text, ["qualitative", "quantitative", "mixed methods"])
        data_type = detect_data_type(text)
        analysis_keywords = ["SPSS", "AMOS", "Stata ", " SEM ", "regression", "correlation", "2SLS", " OLS ", "EViews", "PLS-SEM", "NVivo", "Python", " R ", "Panel Regression", "Logit", "Probit", "Sobel"]
        framework_keywords = [
    # Corporate Governance & Management Theories
         "agency theory", "stewardship theory", "stakeholder theory", "resource dependence theory", "institutional theory", "signaling theory", "upper echelons theory", "transaction cost theory", "managerial hegemony theory", "social contract theory", "political cost theory", "legitimacy theory", "contingency theory", "critical theory", "role theory",

    # Behavioral & Psychological Theories
        "theory of planned behavior", "TPB", "theory of reasoned action", " TRA ", "expectancy theory", "equity theory", "goal-setting theory", "social cognitive theory", "prospect theory", "cognitive dissonance theory", "motivation-hygiene theory",

    # Economic & Financial Theories
        "efficient market hypothesis", "pecking order theory", "trade-off theory", "market timing theory", "modigliani miller theorem", "random walk theory", "portfolio theory", "capital asset pricing model", "CAPM", "arbitrage pricing theory", "option pricing theory", "real options theory",

    # Accounting & Auditing Theories
        "positive accounting theory", "normative accounting theory", "accountability theory", "audit expectation gap theory", "agency cost theory", "accounting conservatism theory", "public interest theory", "capture theory",

    # Strategic & Resource-Based Theories
        "RBV", "resource-based view", "dynamic capabilities theory", "core competence theory", "blue ocean strategy", "disruptive innovation theory", "strategic alignment model",

    # Ethics, CSR, and Sustainability Theories
        "triple bottom line", "stakeholder-agency theory", "carroll's csr pyramid", "corporate social performance theory", "sustainable development theory",

    # Communication & Information Theories
        "media richness theory", "information asymmetry theory", "diffusion of innovations theory",

    # Organizational Theories 
        "open systems theory", "chaos theory", "systems theory", "organizational learning theory", "learning organization theory", "bureaucratic theory", "scientific management theory", "X and Y theory", "path-goal theory", "transformational leadership theory", "servant leadership theory"]
        analysis = find_keywords(text, analysis_keywords)
        frameworks = find_keywords(text, framework_keywords)
        journals = extract_journal_names_dynamic(text)
        years = find_recent_years(text)
        findings = extract_key_findings(text)

        st.header("📊 Evaluation Summary")
        st.markdown("### 🤖 AI Writing Style Check")
        st.write(detect_ai_like_writing(text))

        st.markdown("### 🔍 Methodology & Data Type")
        st.table(pd.DataFrame({"Methodology": methodology if methodology else ["Not Found"], "Data Type": [data_type]}))

        st.markdown("### 📈 Data Analysis Techniques")
        st.write(", ".join(analysis) if analysis else "Not Found")

        st.markdown("### 📚 Theoretical Frameworks")
        st.write(", ".join(frameworks) if frameworks else "Not Found")

        st.markdown("### 📑 Journals Used in Citations")
        if journals:
            st.table(pd.DataFrame(journals, columns=["Journal Name"]))
        else:
            st.write("Not Found")
    if journals != ["Not Found"]:
        st.download_button("⬇️ Download Journals (CSV)", 
                           data=pd.DataFrame(journals, columns=["Journal"]).to_csv(index=False),
                           file_name="journals_list.csv", 
                           mime="text/csv")

        st.markdown("### 📅 Recent References (Last 5 Years)")
        if years:
            st.table(pd.DataFrame(years.items(), columns=["Year", "Count"]).sort_values(by="Year", ascending=False))
        else:
            st.write("No recent references found")

        st.markdown("### 💡 Key Findings / Usefulness of Study")
        for f in findings:
            st.write("- ", f.capitalize())


        summary = {
            "Methodology": methodology,
            "Data Type": data_type,
            "Data Analysis Tools": analysis,
            "Theoretical Frameworks": frameworks,
            "Journals Used": journals,
            "Recent References": years,
            "Key Findings": findings,
        }

        st.markdown("### 📥 Download Evaluation Summary")
        col1, col2 = st.columns(2)

        with col1:
            word_file = generate_docx_report(summary)
            st.download_button("📄 Download Word Document", data=word_file, file_name="Evaluation_Summary.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        with col2:
            pdf_file = generate_pdf_report(summary)
            st.download_button("📄 Download PDF Report", data=pdf_file, file_name="Evaluation_Summary.pdf", mime="application/pdf")


